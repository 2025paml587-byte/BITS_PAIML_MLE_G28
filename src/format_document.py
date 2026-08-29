"""
Format the Project Documentation Word document with improved structure, styles, and layout.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_margins(doc, top=1, bottom=1, left=1, right=1):
    """Set document margins."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def style_heading(paragraph, level, color=(31, 78, 121)):
    """Apply heading style and formatting."""
    paragraph.style = f'Heading {level}'
    for run in paragraph.runs:
        run.font.size = Pt(24 - (level - 1) * 2)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*color)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

def style_body_text(paragraph):
    """Apply consistent body text styling."""
    paragraph.style = 'Normal'
    for run in paragraph.runs:
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15

def style_table(table):
    """Apply consistent table styling."""
    # Set table style
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Style header row
    for cell in table.rows[0].cells:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'E7E6E6')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Style body rows with alternating colors
    for i, row in enumerate(table.rows[1:], 1):
        if i % 2 == 0:
            for cell in row.cells:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F2F2F2')
                cell._element.get_or_add_tcPr().append(shading_elm)
        
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

def format_list_items(paragraph):
    """Format list items with consistent styling."""
    for run in paragraph.runs:
        run.font.size = Pt(11)
    paragraph.paragraph_format.space_after = Pt(4)

def main():
    doc_path = r'docs\Project Documentation_v0_G28.docx'
    doc = Document(doc_path)
    
    # Set document margins
    set_margins(doc, top=1, bottom=1, left=1, right=1)
    
    # Identify and format different elements
    heading_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            para.paragraph_format.space_after = Pt(0)
            continue
        
        # Detect headings by length and content
        # Main heading (single line, short)
        if len(text) < 80 and para.style.name.startswith('Heading'):
            heading_level = int(para.style.name.split()[-1]) if 'Heading' in para.style.name else 1
            style_heading(para, heading_level)
        elif len(text) < 100 and not any(char.isdigit() for char in text[:10]):
            # Could be a heading, apply heading 1 style
            if heading_count < 15:  # Limit to avoid over-styling
                style_heading(para, 1)
                heading_count += 1
            else:
                style_body_text(para)
        else:
            style_body_text(para)
    
    # Format tables
    for table in doc.tables:
        style_table(table)
        # Ensure column widths are reasonable
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_after = Pt(6)
    
    # Add page breaks after major sections (after headings with extra spacing)
    # This will be done by adding spacing before certain headings
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith('Heading 1') and i > 0:
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.page_break_before = False  # Use spacing instead
    
    # Save formatted document
    doc.save(doc_path)
    print(f"✓ Document formatted successfully!")
    print(f"  - Applied consistent heading hierarchy")
    print(f"  - Set standard margins (1 inch)")
    print(f"  - Applied uniform font sizing and spacing")
    print(f"  - Formatted tables with header styling and alternating row colors")
    print(f"  - Enhanced body text with proper line spacing")
    print(f"\nFormatted document saved to: {doc_path}")

if __name__ == '__main__':
    main()
