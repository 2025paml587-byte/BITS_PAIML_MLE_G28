from docx import Document
from pathlib import Path
import re

doc_path = Path('docs/Project Documentation_G28.docx')
doc = Document(str(doc_path))

print("=" * 80)
print("WORD DOCUMENT STRUCTURE - Paragraphs with Placeholders")
print("=" * 80)

placeholder_count = 0
for i, para in enumerate(doc.paragraphs, 1):
    txt = para.text.strip()
    if '[' in txt and (']' in txt or 'Value' in txt or 'Insert' in txt):
        placeholder_count += 1
        print(f"\n[Para {i}] {txt[:200]}")

print("\n" + "=" * 80)
print("TABLES IN DOCUMENT")
print("=" * 80)
for t_idx, table in enumerate(doc.tables, 1):
    print(f"\nTable {t_idx}: ({len(table.rows)} rows x {len(table.columns)} cols)")
    for r_idx, row in enumerate(table.rows[:4], 1):
        row_text = ' | '.join([cell.text[:35] for cell in row.cells])
        print(f"  Row {r_idx}: {row_text}")

print(f"\n\nTotal placeholders found: {placeholder_count}")
print(f"Total tables: {len(doc.tables)}")
